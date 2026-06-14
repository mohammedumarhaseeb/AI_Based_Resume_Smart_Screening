import docx

doc = docx.Document()
doc.add_heading('John Doe', 0)
doc.add_paragraph('Experienced Software Engineer with 5 years of experience in full-stack development.')
doc.add_heading('Skills', level=1)
doc.add_paragraph('Python, Java, React, SQL, AWS, Docker, Kubernetes')
doc.add_paragraph('Communication, Teamwork, Leadership')
doc.save('dummy_resume.docx')
